"""The unaudited financial statements as an editable Word document.

The firm's deliverable: *"A set of unaudited financial statements as an
editable Word document... That document is the deliverable. The app's job
ends when it is produced."*

Editable is the whole point. A PDF is finished; a Word file is where the
preparer changes a figure, rewrites a note, or adds one that was never in
the list. So this does not try to be a perfect rendering - it tries to be a
good starting draft that a person then finishes.

It converts the very HTML the preview already renders, rather than walking
the report model a second time. One rendering path means the Word file and
the on-screen preview cannot drift apart, and every improvement to the
statements shows up in both. The alternative - a separate docx builder
reading the same data - is two things to keep in step, and they never stay
in step.

What survives the conversion: headings, paragraphs, lists, tables with
their header rows, and bold and italic runs. What does not: colour, page
backgrounds, and CSS layout, none of which belong in a document somebody is
about to edit anyway.
"""
import io
import logging
import re
from html.parser import HTMLParser

from docx import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

log = logging.getLogger(__name__)

HEADINGS = {"h1": 0, "h2": 1, "h3": 2, "h4": 3, "h5": 4, "h6": 4}
SKIP = {"script", "style", "head", "nav", "button", "form", "select"}
BLOCKS = {"p", "div", "li", "tr", "section", "article", "header", "footer"}

# A figure column should be right-aligned and a label column should not.
# Deciding by content is the only signal available once the CSS is gone.
NUMERIC = re.compile(r"^[\s(]*-?[\d,]+\.?\d*[\s)%]*$")


class _Reader(HTMLParser):
    """Turn the report's HTML into a flat list of instructions.

    Deliberately not a general HTML renderer. It handles what the report
    templates actually emit and ignores the rest, because a converter that
    tries to handle everything handles nothing predictably.
    """

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.out = []
        self._skip_depth = 0
        self._text = []
        self._bold = 0
        self._italic = 0
        self._heading = None
        self._in_table = False
        self._row = None
        self._cell = None
        self._header_row = False

    # -- text collection --------------------------------------------------

    def _flush(self):
        runs = [r for r in self._text if r[0].strip()]
        self._text = []
        return runs

    def _emit_block(self):
        runs = self._flush()
        if not runs:
            return
        if self._heading is not None:
            self.out.append(("heading", self._heading, runs))
        else:
            self.out.append(("para", None, runs))

    def handle_data(self, data):
        if self._skip_depth:
            return
        if not data.strip():
            # Keep a single space so "<b>Total</b> revenue" does not become
            # "Totalrevenue".
            if self._text and not self._text[-1][0].endswith(" "):
                self._text.append((" ", False, False))
            return
        self._text.append((re.sub(r"\s+", " ", data),
                           bool(self._bold), bool(self._italic)))

    # -- structure --------------------------------------------------------

    def handle_starttag(self, tag, attrs):
        if tag in SKIP:
            self._skip_depth += 1
            return
        if self._skip_depth:
            return

        if tag in ("b", "strong"):
            self._bold += 1
        elif tag in ("i", "em"):
            self._italic += 1
        elif tag == "br":
            self._text.append(("\n", False, False))
        elif tag in HEADINGS:
            self._emit_block()
            self._heading = HEADINGS[tag]
        elif tag == "table":
            self._emit_block()
            self._in_table = True
            self.out.append(("table_start", None, None))
        elif tag == "tr" and self._in_table:
            self._row = []
        elif tag in ("td", "th") and self._in_table:
            self._cell = []
            self._text = []
            if tag == "th":
                self._header_row = True
        elif tag == "hr":
            self._emit_block()
            self.out.append(("rule", None, None))
        elif tag in BLOCKS and not self._in_table:
            self._emit_block()

    def handle_endtag(self, tag):
        if tag in SKIP:
            self._skip_depth = max(0, self._skip_depth - 1)
            return
        if self._skip_depth:
            return

        if tag in ("b", "strong"):
            self._bold = max(0, self._bold - 1)
        elif tag in ("i", "em"):
            self._italic = max(0, self._italic - 1)
        elif tag in HEADINGS:
            self._emit_block()
            self._heading = None
        elif tag in ("td", "th") and self._in_table:
            text = "".join(r[0] for r in self._flush()).strip()
            if self._row is not None:
                self._row.append(text)
            self._cell = None
        elif tag == "tr" and self._in_table:
            if self._row:
                self.out.append(("row", self._header_row, self._row))
            self._row = None
            self._header_row = False
        elif tag == "table":
            self.out.append(("table_end", None, None))
            self._in_table = False
        elif tag == "li":
            runs = self._flush()
            if runs:
                self.out.append(("bullet", None, runs))
        elif tag in BLOCKS and not self._in_table:
            self._emit_block()

    def close(self):
        super().close()
        self._emit_block()
        return self.out


def _write_runs(paragraph, runs):
    for text, bold, italic in runs:
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic


def build(html: str, title: str = None) -> bytes:
    """The report's HTML as a .docx file, returned as bytes."""
    reader = _Reader()
    reader.feed(html)
    instructions = reader.close()

    document = DocxDocument()

    # A statutory set of accounts is set in a serif face at 10 or 11 point.
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)

    if title:
        document.add_heading(title, level=0)

    table = None
    pending_rows = []

    for kind, arg, payload in instructions:
        if kind == "table_start":
            pending_rows = []
            table = True
            continue

        if kind == "row" and table:
            pending_rows.append((bool(arg), payload))
            continue

        if kind == "table_end":
            if pending_rows:
                width = max(len(cells) for _h, cells in pending_rows)
                docx_table = document.add_table(rows=0, cols=width)
                docx_table.style = "Table Grid"
                docx_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for is_header, cells in pending_rows:
                    row = docx_table.add_row()
                    for index in range(width):
                        text = cells[index] if index < len(cells) else ""
                        cell = row.cells[index]
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run(text)
                        run.bold = is_header
                        if NUMERIC.match(text or ""):
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                document.add_paragraph()
            table = None
            pending_rows = []
            continue

        if kind == "heading":
            document.add_heading("", level=min(arg + 1, 4))
            _write_runs(document.paragraphs[-1], payload)
        elif kind == "bullet":
            _write_runs(document.add_paragraph(style="List Bullet"), payload)
        elif kind == "rule":
            document.add_paragraph("_" * 60)
        elif kind == "para":
            _write_runs(document.add_paragraph(), payload)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
