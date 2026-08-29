"""Rule-based extractors: Excel, CSV, Word, typed PDF.

These are deterministic and free. A number read from cell B12 of a spreadsheet
IS that number — no inference involved — which is why they run before the AI
fallback rather than after it.

Each parser returns an `ExtractionResult`; the dispatcher decides whether the
result is good enough or whether Claude should take a second pass.
"""
import csv as csv_module
import io
import logging
import re
from pathlib import Path

from .base import (ExtractedRow, ExtractionResult, clean_label,
                   find_header_row, looks_like_header, parse_amount)

log = logging.getLogger(__name__)


# --------------------------------------------------------------------------
# Column detection — shared by every tabular parser
# --------------------------------------------------------------------------

# A column headed with a year, or with a date carrying one, is the prior-year
# comparative every accounting package prints beside the current figures.
_YEAR = re.compile(r"\b(19|20)\d{2}\b")
COMPARATIVE_WORDS = ("prior year", "previous year", "last year", "comparative",
                     "prior period", "previous period", "prior yr")


def _identify_columns(header_cells):
    """Work out which column holds what, from a header row.

    Returns a dict like {"label": 1, "debit": 3, "credit": 4, "amount": None}.
    Falls back to positional guessing when headers are unhelpful.
    """
    mapping = {"label": None, "debit": None, "credit": None,
               "amount": None, "code": None, "type": None,
               "comparative": set()}

    for idx, cell in enumerate(header_cells):
        text = str(cell or "").strip().lower()
        if not text:
            continue
        # "Dr."/"Cr." are matched as a whole cell, never as a substring:
        # "dr" occurs inside "address" and would mis-tag that column.
        if mapping["debit"] is None and (
                "debit" in text or text.rstrip(".") in ("dr", "dr")):
            mapping["debit"] = idx
        elif mapping["credit"] is None and (
                "credit" in text or text.rstrip(".") == "cr"):
            mapping["credit"] = idx
        elif mapping["code"] is None and any(
                w in text for w in ("code", "a/c", "acct", "account no", "gl")):
            mapping["code"] = idx
        # Tested before "label", because "account type" contains "account"
        # and would otherwise be claimed as the name column.
        elif mapping["type"] is None and text in (
                "type", "account type", "acct type", "a/c type",
                "class", "classification", "category"):
            mapping["type"] = idx
        elif mapping["label"] is None and any(
                w in text for w in ("particular", "description", "account",
                                    "name", "narration", "item")):
            mapping["label"] = idx
        elif mapping["amount"] is None and any(
                w in text for w in ("amount", "balance", "total", "value",
                                    "sgd", "current year")):
            mapping["amount"] = idx

    # A sheet that gives both Dr and Cr usually also carries a derived
    # "Total" column. Reading that as a separate amount would double the
    # figure, so the debit/credit pair wins.
    if mapping["debit"] is not None and mapping["credit"] is not None:
        mapping["amount"] = None

    # Anything left over headed with a year or a date is last year's column,
    # printed beside this one to compare against. A header that names its own
    # side - "DEBIT - YEAR TO DATE" - was claimed above and never reaches
    # here, so only genuinely unclaimed columns can be marked.
    for idx, cell in enumerate(header_cells):
        if idx in (mapping["label"], mapping["debit"], mapping["credit"],
                   mapping["amount"], mapping["code"]):
            continue
        text = str(cell or "").strip().lower()
        if text and (_YEAR.search(text)
                     or any(w in text for w in COMPARATIVE_WORDS)):
            mapping["comparative"].add(idx)

    return mapping


def _row_from_cells(cells, cols, source_ref):
    """Build an ExtractedRow from one spreadsheet/table row."""
    if not cells or all(c is None or str(c).strip() == "" for c in cells):
        return None

    # Find the label: use the mapped column, else the first text-y cell.
    label = ""
    if cols.get("label") is not None and cols["label"] < len(cells):
        label = clean_label(cells[cols["label"]])
    if not label:
        for cell in cells:
            text = clean_label(cell)
            if text and parse_amount(cell) is None:
                label = text
                break

    debit = credit = amount = None
    if cols.get("debit") is not None and cols["debit"] < len(cells):
        debit = parse_amount(cells[cols["debit"]])
    if cols.get("credit") is not None and cols["credit"] < len(cells):
        credit = parse_amount(cells[cols["credit"]])
    if cols.get("amount") is not None and cols["amount"] < len(cells):
        amount = parse_amount(cells[cols["amount"]])

    # No mapped numeric columns? Take the last parseable number on the row —
    # in practice that's the balance column in most exports.
    if debit is None and credit is None and amount is None:
        # Never fall back onto last year's column: it is the rightmost number
        # on the row, so "last number wins" lands on it precisely when this
        # year's cell is blank, and reports a prior-year balance as current.
        # An account code is not an amount either - "4230" is a numeric cell
        # and would be taken as $4,230 on any row with nothing else to find.
        skip = set(cols.get("comparative") or set())
        skip.update(i for i in (cols.get("code"), cols.get("label"),
                                cols.get("type"))
                    if i is not None)
        for idx in range(len(cells) - 1, -1, -1):
            if idx in skip:
                continue
            value = parse_amount(cells[idx])
            if value is not None:
                amount = value
                break

    code = None
    if cols.get("code") is not None and cols["code"] < len(cells):
        code = str(cells[cols["code"]] or "").strip() or None

    account_type = None
    if cols.get("type") is not None and cols["type"] < len(cells):
        account_type = str(cells[cols["type"]] or "").strip()[:60] or None

    # A row with no figure at all is a heading, a title-block line or a
    # spacer - not a line item. Emitting those was filling the review grid
    # with junk the auditor then had to clear by hand. A genuine nil-balance
    # account contributes nothing to the statements either, and can still be
    # added manually in Review & Correct.
    if debit is None and credit is None and amount is None:
        return None

    return ExtractedRow(
        label=label,
        raw_label=label,
        amount=amount,
        debit=debit,
        credit=credit,
        account_code=code,
        account_type=account_type,
        raw_values=[str(c) if c is not None else "" for c in cells],
        source_ref=source_ref,
    )


# --------------------------------------------------------------------------
# Excel
# --------------------------------------------------------------------------

# Sheet names that mean "this is the trial balance", not a schedule.
TB_SHEET_HINTS = ("trial balance", "trialbalance", "tb summary", "trial bal")


def list_sheets(path: Path) -> list:
    """Every sheet in a workbook, with enough detail to choose between them.

    Shown to the auditor so they can say which sheet holds the figures. The
    row and column counts and the first few labels are usually enough to
    tell a trial balance from a transaction ledger at a glance.
    """
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=True, read_only=True)
    sheets = []

    for sheet in workbook.worksheets:
        rows = []
        for index, row in enumerate(sheet.iter_rows(values_only=True)):
            if index >= 40:
                break
            rows.append(list(row))

        header_idx, header_cells = find_header_row(rows)
        labels = []
        start = (header_idx + 1) if header_idx is not None else 0
        for row in rows[start:start + 6]:
            for cell in row:
                text = str(cell or "").strip()
                if text and parse_amount(cell) is None:
                    labels.append(text[:34])
                    break

        filled = sum(1 for row in rows
                     if any(str(c or "").strip() for c in row))
        sheets.append({
            "name": sheet.title,
            "rows": filled,
            "headers": [str(c).strip() for c in (header_cells or [])
                        if str(c or "").strip()][:8],
            "sample": labels[:5],
            "looks_like_tb": any(hint in (sheet.title or "").lower()
                                 for hint in TB_SHEET_HINTS),
        })

    workbook.close()
    return sheets


def _sheets_to_read(workbook, chosen=None):
    """Which sheets in a workbook actually hold the figures.

    A client's "Management Accounts" workbook holds the general ledger, the
    GST computation, a receivables listing, a depreciation schedule and the
    trial balance side by side. Reading all of them merges months of
    transaction detail into the trial balance: dates arrive as account names
    and the totals are meaningless.

    So when a workbook names a trial balance sheet, that is the one read.
    Everything else is a schedule supporting it, already summarised into it,
    and counting it again would double the figures.

    A workbook with no such sheet keeps the old behaviour - every sheet -
    because a plain one-sheet export is still the common case.
    """
    # An explicit choice by the auditor always wins. A sheet named "Trial
    # Balance" is not always a trial balance - one real client workbook uses
    # that name for a cash-movement summary while the audited figures sit on
    # sheets called IS and BS - so the guess below is only a default.
    if chosen:
        wanted = {str(name).strip().lower() for name in chosen}
        picked = [sheet for sheet in workbook.worksheets
                  if (sheet.title or "").strip().lower() in wanted]
        if picked:
            log.info("Reading auditor-chosen sheet(s): %s",
                     [s.title for s in picked])
            return picked

    named = [sheet for sheet in workbook.worksheets
             if any(hint in (sheet.title or "").lower()
                    for hint in TB_SHEET_HINTS)]
    if named:
        log.info("Workbook has %s sheet(s); reading only %s",
                 len(workbook.worksheets), [s.title for s in named])
        return named

    return list(workbook.worksheets)


def _classifier_columns(data_rows, cols):
    """Columns that classify each account rather than carry a figure.

    In a client's trial balance these are headings like Category and FS -
    text, filled in on nearly every real account line. They matter because
    of what they reveal about the rows where they are BLANK.
    """
    numeric = {cols.get(k) for k in ("debit", "credit", "amount", "code")}
    counts, considered = {}, 0

    for row in data_rows[:200]:
        cells = list(row)
        if not any(str(c or "").strip() for c in cells):
            continue
        considered += 1
        for idx, cell in enumerate(cells):
            if idx in numeric:
                continue
            text = str(cell or "").strip()
            if text and parse_amount(cell) is None:
                counts[idx] = counts.get(idx, 0) + 1

    if considered < 5:
        return set()

    # Populated on most lines, so a blank one is a real signal rather than
    # an occasionally empty column.
    frequent = {idx for idx, n in counts.items() if n / considered >= 0.6}

    # The leftmost of these is the account name, not a classification - and
    # a sub-breakdown row DOES carry a name. Keeping it here would mean no
    # row ever looked like a sub-line.
    if cols.get("label") is not None:
        frequent.discard(cols["label"])
    elif frequent:
        frequent.discard(min(frequent))

    return frequent


def _is_subline(cells, classifiers):
    """True for a breakdown row nested under the account above it.

    Client trial balances often expand one account into its parts:

        Bank Charges   Expenses   P&L   76.00
          Bank 1                          50.00
          Bank 2                          26.00

    The child rows carry a figure but no classification, and their amounts
    are already inside the parent's 76.00. Reading them as accounts of their
    own would report 152.00 of bank charges.
    """
    if not classifiers:
        return False
    if not any(parse_amount(c) is not None for c in cells):
        return False
    # Every classifying column blank, yet the row carries a figure.
    return all(not str(cells[idx] or "").strip()
               for idx in classifiers if idx < len(cells))


def extract_xlsx(path: Path, sheets=None) -> ExtractionResult:
    from openpyxl import load_workbook
    from openpyxl.utils import get_column_letter

    result = ExtractionResult(engine="openpyxl")
    # data_only=True gives calculated values rather than formula strings.
    workbook = load_workbook(path, data_only=True, read_only=True)

    for sheet in _sheets_to_read(workbook, chosen=sheets):
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue

        # Find the header row within the first 15 rows; spreadsheets usually
        # carry a title block above the actual table.
        header_idx, header_cells = find_header_row(rows)
        cols, start = {}, 0
        if header_cells is not None:
            cols = _identify_columns(header_cells)
            start = header_idx + 1
        if not cols:
            cols = {"label": None, "debit": None, "credit": None,
                    "amount": None, "code": None}

        classifiers = _classifier_columns(rows[start:], cols)

        for offset, row in enumerate(rows[start:], start=start):
            cells = list(row)

            source_ref = {"sheet": sheet.title, "row": offset + 1}
            if cols.get("amount") is not None:
                source_ref["cell"] = f"{get_column_letter(cols['amount'] + 1)}{offset + 1}"
            extracted = _row_from_cells(cells, cols, source_ref)
            if not extracted:
                continue

            # A row carrying a figure but no classification is EITHER a
            # sub-breakdown already inside the line above it, OR a real
            # account the bookkeeper never classified. Nothing in the file
            # distinguishes the two, and the cost of being wrong is not
            # symmetric: dropping a real account silently loses money from
            # the trial balance, while keeping a sub-line shows a figure
            # twice where the imbalance makes it visible.
            #
            # So it is flagged for the auditor, never decided here.
            if _is_subline(cells, classifiers):
                extracted.needs_review = True
                extracted.confidence = min(extracted.confidence, 0.45)
                extracted.source_ref["unclassified"] = True

            result.rows.append(extracted)

    workbook.close()
    return result


# --------------------------------------------------------------------------
# CSV
# --------------------------------------------------------------------------

def extract_csv(path: Path) -> ExtractionResult:
    result = ExtractionResult(engine="csv")
    raw = path.read_bytes()

    # Try the common encodings before giving up; accounting exports are often
    # Windows-1252 rather than UTF-8.
    text = None
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        result.error = "Could not decode file as text"
        return result

    result.raw_text = text[:20000]

    # Sniff the delimiter — comma, semicolon and tab are all common.
    try:
        dialect = csv_module.Sniffer().sniff(text[:4096], delimiters=",;\t|")
    except csv_module.Error:
        dialect = csv_module.excel

    rows = list(csv_module.reader(io.StringIO(text), dialect))
    if not rows:
        return result

    header_idx, header_cells = find_header_row(rows)
    cols, start = {}, 0
    if header_cells is not None:
        cols = _identify_columns(header_cells)
        start = header_idx + 1
    if not cols:
        cols = {"label": None, "debit": None, "credit": None,
                "amount": None, "code": None}

    for offset, row in enumerate(rows[start:], start=start):
        extracted = _row_from_cells(row, cols, {"row": offset + 1})
        if extracted:
            result.rows.append(extracted)

    return result


# --------------------------------------------------------------------------
# Word
# --------------------------------------------------------------------------

def extract_docx(path: Path) -> ExtractionResult:
    import docx

    result = ExtractionResult(engine="python-docx")
    document = docx.Document(str(path))

    # Tables first — that's where the numbers usually live.
    for table_idx, table in enumerate(document.tables):
        grid = [[cell.text for cell in row.cells] for row in table.rows]
        if not grid:
            continue

        cols = {}
        start = 0
        if looks_like_header(grid[0]):
            cols = _identify_columns(grid[0])
            start = 1
        if not cols:
            cols = {"label": None, "debit": None, "credit": None,
                    "amount": None, "code": None}

        for offset, cells in enumerate(grid[start:], start=start):
            extracted = _row_from_cells(
                cells, cols, {"table": table_idx + 1, "row": offset + 1})
            if extracted:
                result.rows.append(extracted)

    # Keep the body text so the AI fallback has context if tables were empty.
    result.raw_text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    return result


# --------------------------------------------------------------------------
# PDF (typed / text-layer)
# --------------------------------------------------------------------------

def extract_pdf(path: Path) -> ExtractionResult:
    import pdfplumber

    result = ExtractionResult(engine="pdfplumber")
    text_chunks = []

    with pdfplumber.open(path) as pdf:
        result.page_count = len(pdf.pages)

        for page_no, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            text_chunks.append(page_text)

            for table in page.extract_tables() or []:
                if not table:
                    continue
                cols = {}
                start = 0
                if looks_like_header(table[0]):
                    cols = _identify_columns(table[0])
                    start = 1
                if not cols:
                    cols = {"label": None, "debit": None, "credit": None,
                            "amount": None, "code": None}

                for offset, cells in enumerate(table[start:], start=start):
                    extracted = _row_from_cells(
                        cells, cols, {"page": page_no, "row": offset + 1})
                    if extracted:
                        result.rows.append(extracted)

    result.raw_text = "\n".join(text_chunks)

    # A PDF with no extractable text is a scan. Signal that so the dispatcher
    # sends it to Claude, which reads page images directly.
    if not result.raw_text.strip():
        result.error = "scanned"

    return result


# --------------------------------------------------------------------------
# Dispatch by file type
# --------------------------------------------------------------------------

PARSERS = {
    "xlsx": extract_xlsx,
    "xls": extract_xlsx,
    "csv": extract_csv,
    "docx": extract_docx,
    "pdf": extract_pdf,
}


def detect_file_type(filename: str) -> str:
    suffix = Path(filename).suffix.lower().lstrip(".")
    if suffix in ("png", "jpg", "jpeg"):
        return "image"
    return suffix


def run_rule_based(path: Path, file_type: str, sheets=None) -> ExtractionResult:
    """Run the deterministic parser for this file type."""
    parser = PARSERS.get(file_type)
    if parser is None:
        # Images have no rule-based path — they go straight to AI.
        return ExtractionResult(engine="none", error="no rule-based parser")
    try:
        if sheets and parser is extract_xlsx:
            return parser(path, sheets=sheets)
        return parser(path)
    except Exception as exc:                       # noqa: BLE001
        return ExtractionResult(engine=parser.__name__, error=str(exc))
