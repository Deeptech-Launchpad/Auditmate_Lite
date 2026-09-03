"""Demo data that exercises every fix from beta feedback round 2.

Each of the seven points needs something specific to be visible at all, and
the existing demo clients cannot show most of them: a two-column trial
balance, a ledger that both matches and disagrees, a set of signed accounts
with real note wording, a first financial year, a cash flow statement to
file, and a company profile to read.

Real files are written to disk, not just rows to the database. The whole
point is to watch the extraction, the auto-verify rule and the build do their
work - seeded figures that never went through a parser would prove nothing.

    flask seed-feedback            create it
    flask seed-feedback --remove   take it away again
"""
import logging
from datetime import date, datetime
from pathlib import Path

from ..extensions import db
from ..models import Customer, CustomerDocument, Document, FinancialYear
from . import storage

log = logging.getLogger(__name__)

CUSTOMER_NAME = "Harbourfront Marine Services Pte Ltd"

# One account per row: (name, code, current, prior, side)
# Balances in BOTH years, so a real imbalance shows up as a real finding.
TB_ROWS = [
    ("Cash at Bank",              "1010",  45230.00,  38900.00, "Dr"),
    # The firm's own figures from the feedback, so point 6 is checkable
    # against the exact numbers they reported.
    ("Trade Receivables",         "1200", 256134.40, 435463.10, "Dr"),
    ("Plant and Equipment",       "1500", 180000.00, 180000.00, "Dr"),
    ("Accumulated Depreciation",  "1510",  54000.00,  36000.00, "Cr"),
    ("Trade Payables",            "2010",  62140.00,  71200.00, "Cr"),
    ("Bank Loan",                 "2200", 150000.00, 180000.00, "Cr"),
    ("Share Capital",             "3000", 100000.00, 100000.00, "Cr"),
    ("Retained Earnings",         "3100", 115224.40, 267163.10, "Cr"),
    ("Revenue",                   "4000", 612300.00, 540100.00, "Cr"),
    ("Cost of Sales",             "5000", 380000.00, 330000.00, "Dr"),
    ("Staff Costs",               "5100", 145000.00, 130000.00, "Dr"),
    ("Depreciation",              "5200",  18000.00,  18000.00, "Dr"),
    ("Other Expenses",            "5300",  69300.00,  62100.00, "Dr"),
]

# The ledger. Same books, so most of it matches by code and name - but two
# things deliberately do not, and each proves a different half of point 7.
GL_ROWS = [
    ("Cash at Bank",          "1010",  45230.00, "Dr"),
    ("Trade Receivables",     "1200", 256134.40, "Dr"),
    ("Trade Payables",        "2010",  62140.00, "Cr"),
    # Disagrees with the trial balance by 12,300 - an income account, so the
    # comparison is valid and the difference is real.
    ("Revenue",               "4000", 600000.00, "Cr"),
    ("Staff Costs",           "5100", 145000.00, "Dr"),
    # In no trial balance account at all: names itself in the finding.
    ("Director Loan Account", "2300",   9000.00, "Dr"),
]

SIGNED_ACCOUNTS_NOTES = [
    ("1. CORPORATE INFORMATION",
     "Harbourfront Marine Services Pte Ltd is a private company limited by "
     "shares, incorporated and domiciled in Singapore. The principal "
     "activities of the Company are those of marine engineering services and "
     "the repair and maintenance of vessels (SSIC 33150)."),
    ("2. TRADE AND OTHER RECEIVABLES",
     "Trade receivables are non-interest bearing and are generally on 30 to "
     "60 day credit terms granted to established customers. A longer term of "
     "90 days is granted to two shipyard customers under contract."),
    ("3. PROPERTY, PLANT AND EQUIPMENT",
     "Depreciation is calculated on a straight-line basis over the estimated "
     "useful lives of the assets as follows: vessels and marine equipment 10 "
     "years, workshop plant 5 years, office equipment 3 years."),
    ("4. BORROWINGS",
     "The bank loan is secured by a first legal mortgage over the Company's "
     "vessel and is repayable in 60 monthly instalments to December 2028."),
    # No note in the FRS library matches this heading, so it must surface as
    # "NO NOTE" rather than disappearing - which is the case point 3 exists
    # to catch.
    ("5. SHARE-BASED PAYMENTS",
     "The Company operated a share option scheme for senior crew during the "
     "financial year. No options remained outstanding at the year end."),
]

# The company's own first period - 14 May 2023 (incorporation) to 31 Dec
# 2023. A single dated column, nothing else: a real first-year export from
# Xero has no prior period to print beside it, and that absence is exactly
# what proves the point - a two-column reading elsewhere must not invent a
# comparative here because none exists. No Retained Earnings row either;
# opening equity is nil for a company's first period, which this shows by
# leaving the account out rather than by a row reading zero.
FIRST_YEAR_TB_ROWS = [
    ("Cash at Bank",         "1010",  21400.00, "Dr"),
    ("Trade Receivables",    "1200",  22000.00, "Dr"),
    ("Plant and Equipment",  "1500",  45000.00, "Dr"),
    ("Trade Payables",       "2010",  12400.00, "Cr"),
    ("Share Capital",        "3000",  50000.00, "Cr"),
    ("Revenue",              "4000", 180000.00, "Cr"),
    ("Cost of Sales",        "5000", 108000.00, "Dr"),
    ("Staff Costs",          "5100",  38000.00, "Dr"),
    ("Other Expenses",       "5300",   8000.00, "Dr"),
]

ACRA_PROFILE = [
    "ACCOUNTING AND CORPORATE REGULATORY AUTHORITY",
    "BUSINESS PROFILE",
    "",
    "Registration No.: 201934567M",
    "Entity Name: KEPPEL BAY LOGISTICS PTE. LTD.",
    "Entity Type: EXEMPT PRIVATE COMPANY LIMITED BY SHARES",
    "Entity Status: Live Company",
    "Date of Incorporation: 08/07/2019",
    "Registered Office Address: 48 PANDAN ROAD",
    "#03-07 JURONG LOGISTICS HUB",
    "SINGAPORE 609289",
    "Financial Year End: 30/06",
    "",
    "Principal Activities:",
    "Activity (I): FREIGHT FORWARDING SERVICES (SSIC 52291)",
    "Description: SEA FREIGHT FORWARDING AND CUSTOMS BROKERAGE",
    "",
    "Officers / Authorised Representatives:",
    "Name: ONG BOON HUAT   Position: DIRECTOR   Date of Appointment: 08/07/2019",
    "Name: SITI NURHALIZA BINTE OMAR   Position: DIRECTOR   Date of Appointment: 12/01/2021",
    "Name: CHONG WEI LIANG   Position: DIRECTOR   Date of Cessation: 31/03/2023",
    "Name: PINNACLE CORPORATE SERVICES PTE LTD   Position: SECRETARY",
    "",
    "Shareholders:",
    "ONG BOON HUAT   60000 ORDINARY",
    "SITI NURHALIZA BINTE OMAR   40000 ORDINARY",
]


# --------------------------------------------------------------------------
# File builders
# --------------------------------------------------------------------------

def _workbook(path, header, rows):
    import openpyxl
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.append(header)
    for row in rows:
        sheet.append(row)
    workbook.save(path)


def _trial_balance_file(path):
    """Two dated columns, which is what point 6 is about."""
    rows = []
    for name, code, current, prior, side in TB_ROWS:
        rows.append([code, name,
                     current if side == "Dr" else None,
                     current if side == "Cr" else None,
                     prior])
    _workbook(path,
              ["Code", "Account Name", "Debit", "Credit", "31 Dec 2024"],
              rows)


def _first_year_tb_file(path):
    """A first-period trial balance - one dated column, no comparative."""
    rows = [[code, name,
             amount if side == "Dr" else None,
             amount if side == "Cr" else None]
            for name, code, amount, side in FIRST_YEAR_TB_ROWS]
    _workbook(path, ["Code", "Account Name", "Debit", "Credit"], rows)


def _general_ledger_file(path):
    rows = [[code, name,
             amount if side == "Dr" else None,
             amount if side == "Cr" else None]
            for name, code, amount, side in GL_ROWS]
    _workbook(path, ["Code", "Account Name", "Debit", "Credit"], rows)


def _cash_flow_file(path):
    _workbook(path, ["Description", "31 Dec 2025"], [
        ["Cash flows from operating activities", None],
        ["Profit before taxation", 151939.60],
        ["Adjustments for depreciation", 18000.00],
        ["Changes in trade receivables", 179328.70],
        ["Net cash from operating activities", 349268.30],
        ["Cash flows from financing activities", None],
        ["Repayment of bank loan", -30000.00],
        ["Net increase in cash", 6330.00],
    ])


def _docx(path, lines):
    import docx
    document = docx.Document()
    for line in lines:
        document.add_paragraph(line)
    document.save(path)


def _signed_accounts_file(path):
    lines = ["HARBOURFRONT MARINE SERVICES PTE LTD",
             "Notes to the Financial Statements",
             "For the financial year ended 31 December 2024", ""]
    for heading, body in SIGNED_ACCOUNTS_NOTES:
        lines += [heading, body, ""]
    _docx(path, lines)


# --------------------------------------------------------------------------
# Seeding
# --------------------------------------------------------------------------

def _attach(financial_year, filename, category, builder):
    """Write a real file into storage and file it as an unread document."""
    from .extraction import file_sha256
    from .extraction.parsers import detect_file_type

    destination = storage.build_path(financial_year.customer_id,
                                     financial_year.id, filename)
    builder(destination)

    document = Document(
        financial_year_id=financial_year.id,
        original_filename=filename,
        stored_filename=destination.name,
        storage_path=str(destination),
        file_type=detect_file_type(filename),
        size_bytes=destination.stat().st_size,
        sha256=file_sha256(destination),
        category=category,
        category_source="manual",
        # Left unread on purpose. Pressing Analyse is the test.
        extraction_status="queued",
        review_status="pending",
    )
    db.session.add(document)
    return document


def inbox_dir() -> Path:
    """Where the files meant for manual upload are put."""
    directory = Path(storage.storage_root()) / "_demo_inbox"
    directory.mkdir(parents=True, exist_ok=True)
    return directory


def remove() -> bool:
    """Delete the demo customer and everything under it."""
    customer = Customer.query.filter_by(name=CUSTOMER_NAME).first()
    if customer is None:
        return False
    db.session.delete(customer)
    db.session.commit()

    for leftover in inbox_dir().glob("*"):
        try:
            leftover.unlink()
        except OSError:
            pass
    return True


def create(user_id=None) -> dict:
    """Build the demo. Returns a summary of what was made."""
    remove()

    customer = Customer(
        name=CUSTOMER_NAME,
        legal_name=CUSTOMER_NAME,
        entity_type="private_limited",
        uen="201845678K",
        gst_reg_no="M90312345X",
        incorporation_date=date(2018, 5, 14),
        financial_year_end_month=12,
        contact_person="Melissa Tan",
        email="melissa.tan@harbourfrontmarine.example",
        phone="+65 6555 0142",
        directors="Ong Boon Huat\nSiti Nurhaliza Binte Omar",
        company_secretary="Pinnacle Corporate Services Pte Ltd",
        address_line1="12 Harbourfront Place",
        address_line2="#08-03 Marine Tower",
        postal_code="098765",
        country="Singapore",
        books_currency="SGD",
        created_by=user_id,
        engagement_partner_id=user_id,
        notes="Seeded to test beta feedback round 2. Safe to delete.",
    )
    db.session.add(customer)
    db.session.flush()

    # FY2023 exists only to be deleted and to carry the first-year marker,
    # which is what point 5 needs to be visible.
    first_year = FinancialYear(
        customer_id=customer.id, year_label="FY2023",
        start_date=date(2023, 5, 14), end_date=date(2023, 12, 31),
        is_first_year=True, status="in_progress")
    db.session.add(first_year)
    db.session.flush()

    _attach(first_year, "Trial Balance FY2023.xlsx", "trial_balance",
            _first_year_tb_file)

    prior_year = FinancialYear(
        customer_id=customer.id, year_label="FY2024",
        start_date=date(2024, 1, 1), end_date=date(2024, 12, 31),
        status="in_progress")
    db.session.add(prior_year)
    db.session.flush()

    year = FinancialYear(
        customer_id=customer.id, year_label="FY2025",
        start_date=date(2025, 1, 1), end_date=date(2025, 12, 31),
        previous_year_id=prior_year.id, status="in_progress")
    db.session.add(year)
    db.session.flush()

    _attach(year, "Trial Balance FY2025.xlsx", "trial_balance",
            _trial_balance_file)
    _attach(year, "General Ledger FY2025.xlsx", "general_ledger",
            _general_ledger_file)
    _attach(year, "Signed Accounts FY2024.docx", "signed_accounts",
            _signed_accounts_file)

    # Not attached to anything: these are uploaded by hand, because the
    # upload itself is what points 1 and 4 are testing.
    inbox = inbox_dir()
    _cash_flow_file(inbox / "Cash Flow Statement FY2025.xlsx")
    _docx(inbox / "ACRA Business Profile - Keppel Bay Logistics.docx",
          ACRA_PROFILE)

    db.session.commit()

    return {
        "customer_id": customer.id,
        "customer": customer.name,
        "fy2025_id": year.id,
        "fy2024_id": prior_year.id,
        "fy2023_id": first_year.id,
        "inbox": str(inbox),
    }
